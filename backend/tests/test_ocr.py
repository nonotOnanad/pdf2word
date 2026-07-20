import io
import shutil

import docx
import fitz
import pytest

from app.converter import text_pdf_to_docx

HAS_OCR = shutil.which("ocrmypdf") and shutil.which("tesseract")


def test_text_pdf_to_docx_extracts_paragraphs(text_pdf):
    result = text_pdf_to_docx(text_pdf)
    document = docx.Document(io.BytesIO(result))
    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "Hello page 1" in all_text
    assert "Hello page 3" in all_text


@pytest.mark.skipif(not HAS_OCR, reason="ocrmypdf/tesseract not installed")
def test_ocr_adds_text_layer(scanned_image_pdf):
    from app.ocr import ocr_pdf

    searchable = ocr_pdf(scanned_image_pdf)
    doc = fitz.open(stream=searchable, filetype="pdf")
    text = "".join(page.get_text() for page in doc)
    doc.close()
    assert "quick brown fox" in text.lower()


@pytest.mark.skipif(not HAS_OCR, reason="ocrmypdf/tesseract not installed")
def test_scanned_pdf_end_to_end(scanned_image_pdf):
    """Full pipeline: scanned image PDF -> OCR -> editable docx with the text."""
    from app.ocr import ocr_pdf

    docx_bytes = text_pdf_to_docx(ocr_pdf(scanned_image_pdf))
    document = docx.Document(io.BytesIO(docx_bytes))
    all_text = " ".join(p.text for p in document.paragraphs).lower()
    assert "quick brown fox" in all_text
    assert "12345" in all_text

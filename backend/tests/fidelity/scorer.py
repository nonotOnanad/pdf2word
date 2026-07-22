"""Reference-free fidelity scoring for PDF -> DOCX conversions.

Given the source PDF bytes and the produced DOCX bytes, compute quality
metrics WITHOUT needing a hand-made "correct" answer. The point is to catch
regressions and compare settings/tools objectively:

  - text_retention : fraction of the PDF's words that survived into the DOCX
  - length_ratio   : DOCX text length / PDF text length (sanity band)
  - tables_found   : number of tables detected in the DOCX
  - table_cells    : total table cells in the DOCX
  - next_column    : True if a cross-app-breaking `nextColumn` section survived
  - valid_docx     : DOCX opened without error
  - paragraphs     : top-level paragraph count

`composite` is a 0-100 roll-up for quick ranking; the individual metrics are
what you actually diagnose with.
"""
from __future__ import annotations

import io
import re
import zipfile
from collections import Counter
from dataclasses import dataclass, asdict

import fitz  # PyMuPDF
import docx  # python-docx

_WORD = re.compile(r"[0-9a-z]+")


def _tokens(text: str) -> Counter:
    return Counter(_WORD.findall(text.lower()))


def pdf_text(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def docx_text(document: "docx.document.Document") -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _has_next_column(docx_bytes: bytes) -> bool:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    return 'w:val="nextColumn"' in xml


@dataclass
class Fidelity:
    valid_docx: bool
    text_retention: float
    length_ratio: float
    tables_found: int
    table_cells: int
    next_column: bool
    paragraphs: int
    composite: float

    def as_dict(self) -> dict:
        return asdict(self)


def score(pdf_bytes: bytes, docx_bytes: bytes) -> Fidelity:
    src = pdf_text(pdf_bytes)
    src_tokens = _tokens(src)

    try:
        document = docx.Document(io.BytesIO(docx_bytes))
        valid = True
    except Exception:
        # Invalid DOCX is a hard failure: nothing else is meaningful.
        return Fidelity(False, 0.0, 0.0, 0, 0, False, 0, 0.0)

    out = docx_text(document)
    out_tokens = _tokens(out)

    # token recall (multiset): how many source word-occurrences survived
    matched = sum((src_tokens & out_tokens).values())
    total = sum(src_tokens.values()) or 1
    text_retention = matched / total

    src_len = len(src.strip()) or 1
    length_ratio = len(out.strip()) / src_len

    tables = document.tables
    cells = sum(len(r.cells) for t in tables for r in t.rows)
    next_col = _has_next_column(docx_bytes)

    # composite: text retention dominates; penalize the cross-app column bug
    # and wildly-off length. Tables are reported, not scored (reference-free
    # scoring can't know how many tables the PDF "should" have).
    length_ok = 1.0 if 0.5 <= length_ratio <= 2.0 else max(0.0, 1 - abs(length_ratio - 1))
    composite = 100 * (0.75 * text_retention + 0.15 * (0.0 if next_col else 1.0) + 0.10 * length_ok)

    return Fidelity(
        valid_docx=valid,
        text_retention=round(text_retention, 4),
        length_ratio=round(length_ratio, 3),
        tables_found=len(tables),
        table_cells=cells,
        next_column=next_col,
        paragraphs=len(document.paragraphs),
        composite=round(composite, 1),
    )

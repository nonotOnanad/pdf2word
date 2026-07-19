import io
import os
import tempfile

import docx  # python-docx, test-only dependency

from app.converter import convert_pdf_to_docx


def test_converts_text_pdf_to_valid_docx(text_pdf):
    result = convert_pdf_to_docx(text_pdf)
    document = docx.Document(io.BytesIO(result))  # raises if not a valid docx
    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "Hello page 1" in all_text


def test_converts_table_pdf(table_pdf):
    result = convert_pdf_to_docx(table_pdf)
    document = docx.Document(io.BytesIO(result))
    assert len(result) > 0
    # content lands either in paragraphs or a detected table
    all_text = "\n".join(p.text for p in document.paragraphs)
    table_text = "".join(
        cell.text for t in document.tables for row in t.rows for cell in row.cells
    )
    assert "Widgets" in (all_text + table_text)


def test_no_temp_files_left_behind(text_pdf):
    tmp = tempfile.gettempdir()
    before = set(os.listdir(tmp))
    convert_pdf_to_docx(text_pdf)
    after = set(os.listdir(tmp))
    leftovers = [f for f in after - before if "pdf2word" in f]
    assert leftovers == []
